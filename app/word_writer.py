# app/word_writer.py

from docx import Document
from datetime import datetime
import os


def write_report_to_word(title: str, content: str, filename: str | None = None):
    doc = Document()

    # 标题
    doc.add_heading(title, 0)

    # 时间
    doc.add_paragraph(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # 分割内容（按段落）
    paragraphs = content.split("\n")

    for p in paragraphs:
        if p.strip():
            doc.add_paragraph(p)

    # 文件名
    if not filename:
        filename = f"ai_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    # 保存路径（项目根目录 /reports）
    output_dir = "reports"
    os.makedirs(output_dir, exist_ok=True)

    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)

    return file_path